from __future__ import annotations

import ast
import json
from collections import Counter
from pathlib import Path

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
DOCX_PATH = REPORT_DIR / "Football_Analysis_Project_Report.docx"
PDF_PATH = REPORT_DIR / "Football_Analysis_Project_Report.pdf"

SOURCE_FILES = [
    "main.py",
    "yolo_inference.py",
    "trackers/tracker.py",
    "trackers/__init__.py",
    "team_assigner/team_assigner.py",
    "team_assigner/__init__.py",
    "player_ball_assigner/player_ball_assigner.py",
    "player_ball_assigner/__init__.py",
    "camera_movement_estimator/camera_movement_estimator.py",
    "camera_movement_estimator/__init__.py",
    "speed_and_distance_estimator/speed_and_distance_estimator.py",
    "speed_and_distance_estimator/__init__.py",
    "view_transformer/view_transformer.py",
    "view_transformer/__int__.py",
    "utils/bbox_utils.py",
    "utils/video_utils.py",
    "utils/__init__.py",
    "README.md",
    "requirements.txt",
    "training/football_training_yolo_v5.ipynb",
    "development_and_analysis/color_assignment.ipynb",
    "training/football-players-detection-1/data.yaml",
    "training/football-players-detection-1/README.dataset.txt",
    "training/football-players-detection-1/README.roboflow.txt",
]


def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(errors="ignore")


def extract_py_metadata(path: Path) -> dict:
    source = read_text(path)
    tree = ast.parse(source)
    imports = []
    functions = []
    classes = []

    for node in tree.body:
        if isinstance(node, ast.Import):
            imports.extend(alias.name for alias in node.names)
        elif isinstance(node, ast.ImportFrom):
            module = node.module or ""
            imports.append(f"{module}: {', '.join(alias.name for alias in node.names)}")
        elif isinstance(node, ast.FunctionDef):
            functions.append({"name": node.name, "line": node.lineno})
        elif isinstance(node, ast.ClassDef):
            methods = []
            for child in node.body:
                if isinstance(child, ast.FunctionDef):
                    methods.append({"name": child.name, "line": child.lineno})
            classes.append({"name": node.name, "line": node.lineno, "methods": methods})

    return {
        "imports": imports,
        "functions": functions,
        "classes": classes,
    }


def extract_notebook_summary(path: Path) -> dict:
    notebook = json.loads(read_text(path))
    previews = []
    for index, cell in enumerate(notebook.get("cells", [])):
        source = "".join(cell.get("source", [])).strip()
        if source:
            previews.append({
                "index": index,
                "type": cell.get("cell_type", ""),
                "source": source,
            })
    return {
        "cell_count": len(notebook.get("cells", [])),
        "previews": previews,
    }


def parse_main_runtime() -> dict:
    source = read_text(ROOT / "main.py")
    video_path = "unknown"
    resize_width = "default"
    max_frames = "full"
    track_stub = "unknown"
    camera_stub = "unknown"
    output_path = "unknown"

    for line in source.splitlines():
        stripped = line.strip()
        if stripped.startswith("input_video_path ="):
            video_path = stripped.split("=", 1)[1].strip().strip("'\"")
        elif "read_video(" in stripped:
            if "resize_width=" in stripped:
                resize_width = stripped.split("resize_width=", 1)[1].split(",", 1)[0].rstrip(")")
            if "max_frames=" in stripped:
                max_frames = stripped.split("max_frames=", 1)[1].split(")", 1)[0]
        elif "stub_path=f'stubs/" in stripped and "track" in stripped:
            track_stub = stripped.split("stub_path=", 1)[1].rstrip(")")
        elif "stub_path=f'stubs/" in stripped and "camera" in stripped:
            camera_stub = stripped.split("stub_path=", 1)[1].rstrip(")")
        elif "save_video(" in stripped:
            output_path = stripped.split(",", 1)[1].strip().rstrip(")")

    return {
        "video_path": video_path,
        "resize_width": resize_width,
        "max_frames": max_frames,
        "track_stub": track_stub,
        "camera_stub": camera_stub,
        "output_path": output_path,
    }


def dataset_counts() -> dict:
    base = ROOT / "training" / "football-players-detection-1" / "football-players-detection-1"
    counts = {}
    for split in ("train", "valid", "test"):
        image_count = len(list((base / split / "images").glob("*")))
        label_count = len(list((base / split / "labels").glob("*")))
        counts[split] = {"images": image_count, "labels": label_count}
    return counts


def tree_lines() -> list[str]:
    return [
        "Football_Analysis/",
        "  main.py",
        "  yolo_inference.py",
        "  requirements.txt",
        "  README.md",
        "  models/",
        "    best.pt",
        "  input_videos/",
        "    08fd33_4.mp4, 1.mp4, 2.mp4, 3.mp4, 13433649_3840_2160_30fps.mp4, fbb.mp4",
        "  output_videos/",
        "    *_output.avi rendered analysis videos",
        "  stubs/",
        "    per-video cached track and camera movement pickle files",
        "  trackers/",
        "    tracker.py, __init__.py",
        "  team_assigner/",
        "    team_assigner.py, __init__.py",
        "  player_ball_assigner/",
        "    player_ball_assigner.py, __init__.py",
        "  camera_movement_estimator/",
        "    camera_movement_estimator.py, __init__.py",
        "  speed_and_distance_estimator/",
        "    speed_and_distance_estimator.py, __init__.py",
        "  view_transformer/",
        "    view_transformer.py, __int__.py",
        "  utils/",
        "    bbox_utils.py, video_utils.py, __init__.py",
        "  training/",
        "    football_training_yolo_v5.ipynb",
        "    football-players-detection-1/",
        "      data.yaml, README.dataset.txt, README.roboflow.txt",
        "      football-players-detection-1/{train,valid,test}/{images,labels}",
        "  development_and_analysis/",
        "    color_assignment.ipynb",
        "  reports/",
        "    build_project_report.py, generated report artifacts",
    ]


def build_context() -> dict:
    py_meta = {}
    all_imports = Counter()
    for rel in SOURCE_FILES:
        path = ROOT / rel
        if path.suffix == ".py":
            meta = extract_py_metadata(path)
            py_meta[rel] = meta
            all_imports.update(meta["imports"])

    training_nb = extract_notebook_summary(ROOT / "training" / "football_training_yolo_v5.ipynb")
    color_nb = extract_notebook_summary(ROOT / "development_and_analysis" / "color_assignment.ipynb")
    data_yaml = yaml.safe_load(read_text(ROOT / "training" / "football-players-detection-1" / "data.yaml"))
    runtime = parse_main_runtime()

    return {
        "py_meta": py_meta,
        "training_nb": training_nb,
        "color_nb": color_nb,
        "data_yaml": data_yaml,
        "dataset_counts": dataset_counts(),
        "dataset_readme": read_text(ROOT / "training" / "football-players-detection-1" / "README.dataset.txt").strip(),
        "dataset_roboflow": read_text(ROOT / "training" / "football-players-detection-1" / "README.roboflow.txt").strip(),
        "requirements": [line.strip() for line in read_text(ROOT / "requirements.txt").splitlines() if line.strip()],
        "runtime": runtime,
        "imports": sorted(all_imports),
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 24, RGBColor(26, 54, 93)),
        ("Heading 1", 16, RGBColor(26, 54, 93)),
        ("Heading 2", 12.5, RGBColor(54, 90, 140)),
        ("Heading 3", 11, RGBColor(60, 60, 60)),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Football Analysis Project Report")
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(26, 54, 93)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Comprehensive implementation, training, dataset, file, function, "
        "input/output, and library reference for the Football_Analysis repository."
    ).italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Generated from the current local repository state in ").bold = True
    meta.add_run(str(ROOT))

    doc.add_paragraph("")
    card = doc.add_table(rows=4, cols=2)
    card.alignment = WD_TABLE_ALIGNMENT.CENTER
    card.style = "Table Grid"
    fields = [
        ("Primary entry point", "main.py"),
        ("Model artifact", "models/best.pt"),
        ("Training notebook", "training/football_training_yolo_v5.ipynb"),
        ("Dataset format", "YOLOv5 / Ultralytics detection labels"),
    ]
    for idx, (left, right) in enumerate(fields):
        row = card.rows[idx].cells
        row[0].text = left
        row[1].text = right
        set_cell_shading(row[0], "DCE6F1")

    doc.add_page_break()


def add_bullets(doc: Document, items: list[str], level: int = 0) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        set_cell_shading(hdr_cells[i], "DCE6F1")

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)

    doc.add_paragraph("")


def compose_docx(context: dict) -> None:
    doc = Document()
    style_document(doc)
    add_cover(doc)

    doc.add_heading("1. Project Overview", level=1)
    doc.add_paragraph(
        "This repository implements a football video analytics pipeline that detects "
        "players, referees, goalkeepers, and the ball; tracks those detections over time; "
        "assigns teams by jersey color; estimates camera motion; projects player movement into "
        "a field coordinate system; and computes ball control, speed, and distance overlays."
    )
    add_bullets(doc, [
        "Primary runtime entry point: main.py",
        "Current configured input video: " + context["runtime"]["video_path"],
        "Current frame read cap: " + str(context["runtime"]["max_frames"]),
        "Current resize setting: " + str(context["runtime"]["resize_width"]),
        "Output video target: " + context["runtime"]["output_path"],
    ])

    doc.add_heading("2. Repository Structure", level=1)
    doc.add_paragraph("The following simplified tree lists the important project files and folders.")
    tree_para = doc.add_paragraph()
    tree_para.style = "No Spacing"
    for line in tree_lines():
        tree_para.add_run(line + "\n")

    doc.add_heading("3. How the Project Runs", level=1)
    add_table(doc, ["Stage", "What happens", "Main code location"], [
        ["Video ingestion", "Reads frames from the selected input video into memory with optional resizing and frame limiting.", "utils/video_utils.py -> read_video"],
        ["Object detection", "Loads models/best.pt with Ultralytics YOLO and predicts detections in batches of 20 frames.", "trackers/tracker.py -> Tracker.detect_frames"],
        ["Tracking", "Converts detections into supervision format and feeds them to ByteTrack.", "trackers/tracker.py -> Tracker.get_object_tracks"],
        ["Position extraction", "Stores player foot positions and ball centers per frame.", "trackers/tracker.py -> Tracker.add_position_to_tracks"],
        ["Camera motion", "Uses Lucas-Kanade optical flow on edge features to estimate frame-to-frame camera motion.", "camera_movement_estimator/camera_movement_estimator.py"],
        ["View transform", "Projects adjusted image points into a simplified field coordinate system.", "view_transformer/view_transformer.py"],
        ["Ball interpolation", "Fills missing ball boxes by pandas interpolation and backfill.", "trackers/tracker.py -> Tracker.interpolate_ball_positions"],
        ["Speed and distance", "Measures transformed displacement over a sliding frame window and converts to km/h.", "speed_and_distance_estimator/speed_and_distance_estimator.py"],
        ["Team assignment", "Uses KMeans on top-half jersey crops to cluster players into two teams.", "team_assigner/team_assigner.py"],
        ["Ball possession", "Assigns the ball to the nearest eligible player and accumulates team control over frames.", "player_ball_assigner/player_ball_assigner.py + main.py"],
        ["Rendering", "Draws ellipses, triangles, ball control percentages, camera motion, and speed/distance text.", "trackers/tracker.py + camera_movement_estimator + speed_and_distance_estimator"],
        ["Output write", "Writes the rendered frames to an AVI file.", "utils/video_utils.py -> save_video"],
    ], widths=[1.3, 4.3, 2.1])

    doc.add_heading("4. Input, Output, and Cached Artifacts", level=1)
    add_table(doc, ["Artifact type", "Examples", "Purpose"], [
        ["Input videos", "input_videos/08fd33_4.mp4, 1.mp4, 2.mp4, 3.mp4, fbb.mp4", "Source match clips consumed by main.py or quick inference scripts."],
        ["Model weights", "models/best.pt", "Custom trained detector used by the main pipeline."],
        ["Cached stubs", "stubs/08fd33_4_track_stubs.pkl, stubs/1_camera_movement_stub.pkl", "Pickled tracking and camera motion results to avoid recomputation."],
        ["Output videos", "output_videos/08fd33_4_output.avi, 1_output.avi", "Rendered analysis videos with overlays and statistics."],
        ["Notebook outputs", "Colab copies of best.pt / last.pt per training notebook cells", "Saved training artifacts after notebook execution."],
    ], widths=[1.4, 3.0, 3.3])

    doc.add_heading("5. Training Workflow and Dataset Usage", level=1)
    doc.add_paragraph(
        "The training notebook training/football_training_yolo_v5.ipynb documents a short Ultralytics training flow. "
        "It installs ultralytics and roboflow, downloads dataset version 1 from Roboflow in YOLOv5 format, rearranges "
        "the train/test/valid directories, launches training, and copies best.pt and last.pt to Google Drive."
    )
    add_bullets(doc, [
        "Roboflow workspace/project from notebook: roboflow-jvuqo / football-players-detection-3zvbc",
        "Download format: yolov5",
        "Notebook training command: yolo task=detect mode=train model=yolov5x.pt data={dataset.location}/data.yaml epochs=10 imgsz=640",
        "Saved artifacts in notebook: runs/detect/train/weights/best.pt and last.pt",
    ])
    doc.add_paragraph("Dataset metadata from local files:")
    add_table(doc, ["Dataset property", "Value"], [
        ["Dataset name", "football-players-detection - v1"],
        ["Source URL", context["data_yaml"]["roboflow"]["url"]],
        ["License", context["data_yaml"]["roboflow"]["license"]],
        ["Classes", ", ".join(context["data_yaml"]["names"])],
        ["Image count noted by Roboflow", "663 images"],
        ["Train split count", str(context["dataset_counts"]["train"]["images"])],
        ["Valid split count", str(context["dataset_counts"]["valid"]["images"])],
        ["Test split count", str(context["dataset_counts"]["test"]["images"])],
        ["Augmentations", "50% horizontal flip; brightness adjustment between -20% and +20%"],
    ], widths=[2.5, 4.8])

    doc.add_heading("6. Notebook and Analysis File Summary", level=1)
    add_table(doc, ["Notebook/file", "Role", "Important contents"], [
        ["training/football_training_yolo_v5.ipynb", "Training and model export", "Package install, Roboflow download, directory moves, yolo detect train command, Drive export of best.pt and last.pt."],
        ["development_and_analysis/color_assignment.ipynb", "Prototype for jersey clustering", "Loads a cropped player image, isolates the upper half, runs 2-cluster KMeans, and inspects corner clusters to isolate player color."],
        ["README.md", "Human-readable overview", "Project setup, dependencies, methodology, and a high-level feature list."],
        ["training/.../data.yaml", "Training configuration", "4 classes: ball, goalkeeper, player, referee; train/val/test paths."],
        ["training/.../README.dataset.txt", "Dataset export notes", "Roboflow export date, image count, annotation format, and augmentations."],
        ["training/.../README.roboflow.txt", "Dataset provenance", "Dataset name, universe URL, provider, and CC BY 4.0 license."],
    ], widths=[2.5, 1.8, 3.0])

    doc.add_heading("7. Function and Class Catalog", level=1)
    doc.add_paragraph(
        "This section enumerates the declared Python functions, classes, and methods present in the repository."
    )
    for rel_path in [p for p in SOURCE_FILES if p.endswith(".py")]:
        meta = context["py_meta"].get(rel_path, {})
        doc.add_heading(rel_path, level=2)
        if meta.get("imports"):
            doc.add_paragraph("Imports: " + "; ".join(meta["imports"]))
        if meta.get("functions"):
            rows = [[item["name"], str(item["line"]), "Top-level function"] for item in meta["functions"]]
            add_table(doc, ["Function", "Line", "Notes"], rows, widths=[2.2, 0.8, 4.5])
        if meta.get("classes"):
            for cls in meta["classes"]:
                doc.add_paragraph(f"Class {cls['name']} (declared at line {cls['line']})")
                rows = [[m["name"], str(m["line"]), f"Method of {cls['name']}"] for m in cls["methods"]]
                add_table(doc, ["Method", "Line", "Notes"], rows, widths=[2.2, 0.8, 4.5])
        if not meta.get("functions") and not meta.get("classes"):
            doc.add_paragraph("This file primarily re-exports symbols or runs simple script code without declared functions.")

    doc.add_heading("8. Module-by-Module Logic Explanation", level=1)
    explanations = [
        ("main.py", [
            "Defines the project orchestration function main().",
            "Selects an input video, resolves a per-video stub/output naming scheme, reads frames, and invokes each processing stage in order.",
            "Stores team information inside tracks['players'][frame_num][player_id]['team'] as a dictionary with id and team_color.",
            "Creates team_ball_control as a NumPy array so downstream overlay logic can use boolean masks.",
        ]),
        ("trackers/tracker.py", [
            "Creates the Ultralytics YOLO model and supervision ByteTrack tracker.",
            "Tracker.get_object_tracks() is the central detection-to-structured-tracks conversion method.",
            "Goalkeepers are mapped into the player class for simpler downstream team handling.",
            "The module also owns most video drawing logic: player ellipses, ball triangles, and ball-control overlay text.",
        ]),
        ("team_assigner/team_assigner.py", [
            "Uses 2-cluster KMeans on the upper half of a player's crop to estimate jersey color.",
            "assign_team_color() builds a team-level color model using the first frame's player detections.",
            "get_player_team() caches assigned team ids by player id to avoid recomputation each frame.",
            "Two hard-coded overrides exist: player id 156 is forced to team 1 and player id 510 is forced to team 2.",
        ]),
        ("player_ball_assigner/player_ball_assigner.py", [
            "Treats the ball owner as the player whose left or right foot point is closest to the ball center.",
            "Only players within max_player_ball_distance = 70 pixels are eligible.",
            "Returns -1 if no player is close enough, which main.py interprets as keep previous team possession or zero at startup.",
        ]),
        ("camera_movement_estimator/camera_movement_estimator.py", [
            "Extracts corner features from left and right edge strips, then runs Lucas-Kanade optical flow across frames.",
            "Stores per-frame x/y camera displacement so track positions can be motion compensated.",
            "Recent robustness updates reacquire features if no valid corners are available for a frame.",
        ]),
        ("view_transformer/view_transformer.py", [
            "Defines a four-point perspective transform from image coordinates to a simplified pitch plane.",
            "Only points inside the configured polygon are transformed; outside points return None.",
            "The transformed coordinates are later used for speed and distance computation.",
        ]),
        ("speed_and_distance_estimator/speed_and_distance_estimator.py", [
            "Computes displacement over frame windows of five frames and uses frame_rate = 24 to convert displacement into speed.",
            "Formula in code: speed_km_per_hour = (distance_covered / time_elapsed) * 3.6.",
            "Also accumulates total distance per tracked player and overlays both values under the player marker.",
        ]),
        ("utils/bbox_utils.py and utils/video_utils.py", [
            "Provide common geometry helpers for center, width, distance, and foot-position calculations.",
            "read_video() supports optional resize_width and max_frames, while save_video() writes AVI output with XVID codec.",
        ]),
        ("yolo_inference.py", [
            "A standalone inference script that loads either yolov8m.pt or models/best.pt and predicts directly on 08fd33_4.mp4.",
            "Useful for quick detector sanity checks outside the larger tracking pipeline.",
        ]),
    ]
    for heading, bullets in explanations:
        doc.add_heading(heading, level=2)
        add_bullets(doc, bullets)

    doc.add_heading("9. Calculation Details", level=1)
    add_table(doc, ["Metric or output", "How it is computed"], [
        ["Ball position", "Center point of the detected ball bounding box using get_center_of_bbox()."],
        ["Player position", "Foot point of the bounding box, defined as ((x1+x2)/2, y2)."],
        ["Ball interpolation", "Missing ball bounding boxes are interpolated with pandas DataFrame.interpolate() and backfilled."],
        ["Team colors", "For each player crop, a 2-cluster KMeans is run on the upper half of the image; the dominant corner cluster is treated as background and the other cluster center is treated as jersey color."],
        ["Team id", "Player jersey color is predicted through a second KMeans model fit over initial player colors, then shifted from 0/1 to 1/2."],
        ["Ball owner", "Nearest player foot endpoint to the ball center, if within 70 pixels."],
        ["Ball control percentage", "Count frames labeled team 1 or team 2 up to the current frame and divide by the total labeled team frames."],
        ["Camera movement", "Largest Lucas-Kanade tracked feature displacement among selected edge features between consecutive frames."],
        ["Perspective transform", "cv2.getPerspectiveTransform() maps configured image polygon vertices into target field coordinates."],
        ["Speed", "distance_covered / time_elapsed, then multiplied by 3.6 to convert m/s to km/h."],
        ["Distance covered", "Running sum of transformed displacement per tracked player across frame windows."],
    ], widths=[2.1, 5.2])

    doc.add_heading("10. Console Output and Rendered Video Output Explained", level=1)
    add_bullets(doc, [
        "Lines like '384x640 15 players, 1 referee, 402.1ms' come from Ultralytics prediction logging. They summarize the resized inference tensor shape, detected classes in that frame, and inference time.",
        "Repeated dictionaries like {0: 'ball', 1: 'goalkeeper', 2: 'player', 3: 'referee'} are printed by Tracker.get_object_tracks() for each frame because the code prints detection.names.",
        "The final AVI output contains player ellipses, player ids, a red triangle above the player currently judged to have the ball, a green triangle above the ball, team ball control percentages, camera motion text, and speed/distance text for tracked players.",
        "Stub pickle files are optimization artifacts only. They do not contain rendered frames; they contain cached structured track dictionaries and camera-motion vectors.",
    ])

    doc.add_heading("11. Useful Libraries and What They Do Here", level=1)
    add_table(doc, ["Library", "Used in project for"], [
        ["ultralytics", "YOLO model loading, training command compatibility, and batched object detection inference."],
        ["supervision", "ByteTrack integration and conversion from Ultralytics results to trackable detections."],
        ["opencv-python (cv2)", "Video I/O, drawing overlays, color conversion, optical flow, perspective transforms, and text rendering."],
        ["numpy", "Array storage, point math, boolean masking, perspective point formatting, and control arrays."],
        ["pandas", "Interpolation and backfilling of missing ball detections."],
        ["scikit-learn", "KMeans clustering for jersey-color grouping and team assignment."],
        ["scipy", "Imported in tracker.py, though scipy.stats.triang is not actively used in the current code."],
        ["pickle", "Reading and writing stub caches for faster repeated runs."],
        ["yaml / PyYAML", "Reading dataset metadata and training configuration files such as data.yaml."],
        ["roboflow", "Used only in the training notebook to download the dataset from Roboflow."],
    ], widths=[1.8, 5.5])

    doc.add_heading("12. Known Constraints and Observed Issues", level=1)
    add_bullets(doc, [
        "The runtime pipeline stores all frames in memory, so full-length high-resolution videos can still become memory-heavy.",
        "The project currently writes AVI output with a fixed 24 FPS in save_video(), even if the source video used a different frame rate.",
        "Track and camera-motion stubs must match the selected input video; otherwise frame-count mismatches can trigger index errors.",
        "The supervision ByteTrack API currently emits a deprecation warning in newer supervision versions.",
        "view_transformer/__int__.py appears to be a typo for __init__.py, though the code currently imports ViewTransformer directly from view_transformer.py.",
        "README.md contains a conceptual movie recommendation section that does not match the football-analysis codebase and should be treated as unrelated noise.",
    ])

    doc.add_heading("13. Source File Inventory Included in this Report", level=1)
    rows = [[rel, "Present" if (ROOT / rel).exists() else "Missing"] for rel in SOURCE_FILES]
    add_table(doc, ["File", "Status"], rows, widths=[5.5, 1.5])

    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ReportTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=22,
        textColor=colors.HexColor("#1A365D"),
        alignment=TA_CENTER,
        spaceAfter=18,
    ))
    styles.add(ParagraphStyle(
        name="Section",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=15,
        textColor=colors.HexColor("#1A365D"),
        spaceBefore=12,
        spaceAfter=8,
    ))
    styles.add(ParagraphStyle(
        name="SubSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=11.5,
        textColor=colors.HexColor("#365A8C"),
        spaceBefore=8,
        spaceAfter=6,
    ))
    styles["BodyText"].fontName = "Helvetica"
    styles["BodyText"].fontSize = 9.5
    styles["BodyText"].leading = 12
    return styles


def pdf_table(data, widths):
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCE6F1")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1A365D")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 8.5),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return table


def compose_pdf(context: dict) -> None:
    styles = pdf_styles()
    story = []
    story.append(Paragraph("Football Analysis Project Report", styles["ReportTitle"]))
    story.append(Paragraph(
        "Comprehensive implementation, training, dataset, file, function, input/output, and library reference for the Football_Analysis repository.",
        styles["BodyText"],
    ))
    story.append(Spacer(1, 0.4 * cm))

    sections = [
        ("1. Project Overview", [
            "This repository implements a football video analytics pipeline that detects football objects, tracks them over time, assigns teams by jersey color, estimates camera movement, projects movement into field coordinates, and overlays possession, speed, and distance metrics.",
            f"Current input in main.py: {context['runtime']['video_path']}",
            f"Current max frame cap in main.py: {context['runtime']['max_frames']}",
            f"Current output target in main.py: {context['runtime']['output_path']}",
        ]),
        ("2. Dataset and Training Summary", [
            "The training notebook installs ultralytics and roboflow, downloads Roboflow dataset version 1 in YOLOv5 format, reorganizes the split folders, and runs yolo detect train with yolov5x.pt for 10 epochs at imgsz=640.",
            f"Dataset URL: {context['data_yaml']['roboflow']['url']}",
            f"Classes: {', '.join(context['data_yaml']['names'])}",
            f"Split counts: train={context['dataset_counts']['train']['images']}, valid={context['dataset_counts']['valid']['images']}, test={context['dataset_counts']['test']['images']}",
        ]),
        ("3. Core Runtime Flow", [
            "read_video -> Tracker.get_object_tracks -> Tracker.add_position_to_tracks -> CameraMovementEstimator.get_camera_movement -> add_adjust_positions_to_tracks -> ViewTransformer.add_transformed_position_to_tracks -> interpolate_ball_positions -> SpeedAndDistanceEstimator.add_speed_and_distance_to_tracks -> TeamAssigner -> PlayerBallAssigner -> draw annotations -> save_video.",
        ]),
        ("4. Output Explanation", [
            "Ultralytics console lines like '384x640 15 players, 1 referee, 402.1ms' are detector logs, not errors.",
            "Printed dictionaries such as {0: 'ball', 1: 'goalkeeper', 2: 'player', 3: 'referee'} come from tracker.py printing detection.names each frame.",
            "Rendered video output contains player ellipses, ids, triangles for ball and possessor, team-ball-control percentages, camera movement text, and speed/distance text.",
        ]),
    ]
    for title, bullets in sections:
        story.append(Paragraph(title, styles["Section"]))
        for bullet in bullets:
            story.append(Paragraph(f"- {bullet}", styles["BodyText"]))

    story.append(PageBreak())
    story.append(Paragraph("5. Source File Inventory", styles["Section"]))
    inventory_data = [["File", "Status"]] + [[rel, "Present" if (ROOT / rel).exists() else "Missing"] for rel in SOURCE_FILES]
    story.append(pdf_table(inventory_data, [11 * cm, 3.2 * cm]))

    story.append(Paragraph("6. Function and Class Inventory", styles["Section"]))
    for rel_path in [p for p in SOURCE_FILES if p.endswith(".py")]:
        story.append(Paragraph(rel_path, styles["SubSection"]))
        meta = context["py_meta"].get(rel_path, {})
        rows = [["Symbol", "Line", "Kind"]]
        for fn in meta.get("functions", []):
            rows.append([fn["name"], str(fn["line"]), "Function"])
        for cls in meta.get("classes", []):
            rows.append([cls["name"], str(cls["line"]), "Class"])
            for method in cls["methods"]:
                rows.append([f"{cls['name']}.{method['name']}", str(method["line"]), "Method"])
        if len(rows) == 1:
            rows.append(["No declared functions/classes", "-", "Script/re-export"])
        story.append(pdf_table(rows, [9.5 * cm, 2 * cm, 3 * cm]))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=1.4 * cm,
        rightMargin=1.4 * cm,
        topMargin=1.2 * cm,
        bottomMargin=1.2 * cm,
    )
    doc.build(story)


def main() -> None:
    REPORT_DIR.mkdir(exist_ok=True)
    context = build_context()
    compose_docx(context)
    compose_pdf(context)
    print(f"Created {DOCX_PATH}")
    print(f"Created {PDF_PATH}")


if __name__ == "__main__":
    main()
