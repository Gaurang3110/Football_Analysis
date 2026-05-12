from __future__ import annotations

import ast
import json
import math
import textwrap
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image as PILImage, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "reports"
DOCX_PATH = REPORT_DIR / "Football_Analysis_Full_Technical_Report.docx"
PDF_PATH = REPORT_DIR / "Football_Analysis_Full_Technical_Report.pdf"
SUMMARY_PATH = REPORT_DIR / "Football_Analysis_Presentation_Points.md"
HIGHLIGHTS_PATH = REPORT_DIR / "Football_Analysis_Technical_Highlights.md"
DIAGRAM_DIR = REPORT_DIR / "diagrams"

IMPORTANT_FILES = [
    "main.py",
    "yolo_inference.py",
    "README.md",
    "requirements.txt",
    "trackers/tracker.py",
    "trackers/__init__.py",
    "team_assigner/team_assigner.py",
    "team_assigner/__init__.py",
    "player_ball_assigner/player_ball_assigner.py",
    "player_ball_assigner/__init__.py",
    "camera_movement_estimator/camera_movement_estimator.py",
    "camera_movement_estimator/__init__.py",
    "speed_and_distance_estimator/speed_and_distance_estimator.py",
    "speed_and_distance_estimator/__init__.py",
    "view_transformer/view_transformer.py",
    "view_transformer/__int__.py",
    "utils/bbox_utils.py",
    "utils/video_utils.py",
    "utils/__init__.py",
    "training/football_training_yolo_v5.ipynb",
    "development_and_analysis/color_assignment.ipynb",
    "training/football-players-detection-1/data.yaml",
    "training/football-players-detection-1/README.dataset.txt",
    "training/football-players-detection-1/README.roboflow.txt",
]

FOLDER_PURPOSES = {
    "models": "Stores trained model weights used for inference.",
    "input_videos": "Stores source match videos passed to the analytics pipeline.",
    "output_videos": "Stores rendered output videos with overlays and analysis annotations.",
    "stubs": "Stores cached pickle files for tracks and camera movement to avoid repeated recomputation.",
    "trackers": "Contains detection-to-tracking conversion and annotation drawing logic.",
    "team_assigner": "Contains jersey-color clustering logic used to assign players to teams.",
    "player_ball_assigner": "Contains nearest-player logic for assigning ball possession.",
    "camera_movement_estimator": "Contains optical-flow based motion compensation logic.",
    "speed_and_distance_estimator": "Contains transformed-motion speed and distance calculations.",
    "view_transformer": "Contains perspective transform logic for mapping image points to field coordinates.",
    "utils": "Contains shared video and bounding-box helper functions.",
    "training": "Contains model training notebooks, dataset metadata, and local training dataset export.",
    "development_and_analysis": "Contains analysis notebooks used while building or validating parts of the pipeline.",
    "reports": "Contains generated documentation artifacts and report builder scripts.",
}

FILE_PURPOSES = {
    "main.py": "Primary execution script that orchestrates the full inference, tracking, analytics, and rendering pipeline.",
    "yolo_inference.py": "Standalone quick inference script for directly testing YOLO predictions on a video.",
    "trackers/tracker.py": "Core tracking module that runs YOLO inference, structures tracks, interpolates ball detections, and draws overlays.",
    "trackers/__init__.py": "Package export file for Tracker.",
    "team_assigner/team_assigner.py": "Assigns players to one of two teams using KMeans clustering over jersey-color crops.",
    "team_assigner/__init__.py": "Package export file for TeamAssigner.",
    "player_ball_assigner/player_ball_assigner.py": "Chooses the player closest to the detected ball.",
    "player_ball_assigner/__init__.py": "Package export file for PlayerBallAssigner.",
    "camera_movement_estimator/camera_movement_estimator.py": "Computes camera motion using Lucas-Kanade optical flow on edge features and renders motion text.",
    "camera_movement_estimator/__init__.py": "Package export file for CameraMovementEstimator.",
    "speed_and_distance_estimator/speed_and_distance_estimator.py": "Computes and renders player speed and distance using transformed coordinates.",
    "speed_and_distance_estimator/__init__.py": "Package export file for SpeedAndDistanceEstimator.",
    "view_transformer/view_transformer.py": "Defines and applies a perspective transform from image coordinates to pitch coordinates.",
    "view_transformer/__int__.py": "Package export file for ViewTransformer; likely intended to be named __init__.py.",
    "utils/bbox_utils.py": "Geometry helper functions for centers, widths, distances, and player foot locations.",
    "utils/video_utils.py": "Video reading and writing helpers with optional resize and frame limiting.",
    "utils/__init__.py": "Package export file for utility functions.",
    "training/football_training_yolo_v5.ipynb": "Notebook describing dataset download, Ultralytics training, and weight export.",
    "development_and_analysis/color_assignment.ipynb": "Notebook used to prototype jersey-color clustering with KMeans.",
    "training/football-players-detection-1/data.yaml": "Ultralytics training configuration and class metadata file.",
    "training/football-players-detection-1/README.dataset.txt": "Roboflow dataset export metadata and augmentation notes.",
    "training/football-players-detection-1/README.roboflow.txt": "Roboflow dataset provenance and license details.",
    "README.md": "Human-readable overview, setup instructions, and project summary.",
    "requirements.txt": "Python dependency list required to run the project.",
}


@dataclass
class SymbolInfo:
    qualified_name: str
    name: str
    kind: str
    file: str
    line: int
    parameters: list[str] = field(default_factory=list)
    calls: list[str] = field(default_factory=list)


def read_text(path: Path) -> str:
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        return path.read_text(errors="ignore")


def list_project_tree(root: Path, depth: int = 4) -> list[str]:
    lines: list[str] = []

    def walk(path: Path, prefix: str = "", level: int = 0) -> None:
        if level > depth:
            return
        items = sorted([p for p in path.iterdir() if p.name not in {"__pycache__", ".git", ".ultralytics"}], key=lambda p: (p.is_file(), p.name.lower()))
        for index, item in enumerate(items):
            connector = "└── " if index == len(items) - 1 else "├── "
            lines.append(f"{prefix}{connector}{item.name}")
            if item.is_dir():
                extension = "    " if index == len(items) - 1 else "│   "
                walk(item, prefix + extension, level + 1)

    lines.append(root.name + "/")
    walk(root)
    return lines


def get_function_parameters(node: ast.FunctionDef) -> list[str]:
    params = []
    for arg in node.args.args:
        params.append(arg.arg)
    if node.args.vararg:
        params.append("*" + node.args.vararg.arg)
    for arg in node.args.kwonlyargs:
        params.append(arg.arg)
    if node.args.kwarg:
        params.append("**" + node.args.kwarg.arg)
    return params


def collect_call_names(node: ast.AST) -> list[str]:
    calls = []
    for child in ast.walk(node):
        if isinstance(child, ast.Call):
            fn = child.func
            if isinstance(fn, ast.Name):
                calls.append(fn.id)
            elif isinstance(fn, ast.Attribute):
                calls.append(fn.attr)
    return calls


def extract_py_info(path: Path, rel_path: str) -> dict[str, Any]:
    tree = ast.parse(read_text(path))
    imports = []
    symbols: list[SymbolInfo] = []
    source_lines = read_text(path).splitlines()

    for node in tree.body:
        if isinstance(node, ast.Import):
            imports.extend(alias.name for alias in node.names)
        elif isinstance(node, ast.ImportFrom):
            module = node.module or ""
            imports.append(f"{module}: {', '.join(alias.name for alias in node.names)}")
        elif isinstance(node, ast.FunctionDef):
            symbols.append(SymbolInfo(
                qualified_name=node.name,
                name=node.name,
                kind="function",
                file=rel_path,
                line=node.lineno,
                parameters=get_function_parameters(node),
                calls=collect_call_names(node),
            ))
        elif isinstance(node, ast.ClassDef):
            symbols.append(SymbolInfo(
                qualified_name=node.name,
                name=node.name,
                kind="class",
                file=rel_path,
                line=node.lineno,
                parameters=[],
                calls=[],
            ))
            for child in node.body:
                if isinstance(child, ast.FunctionDef):
                    symbols.append(SymbolInfo(
                        qualified_name=f"{node.name}.{child.name}",
                        name=child.name,
                        kind="method",
                        file=rel_path,
                        line=child.lineno,
                        parameters=get_function_parameters(child),
                        calls=collect_call_names(child),
                    ))

    return {
        "imports": imports,
        "symbols": symbols,
        "line_count": len(source_lines),
        "size_bytes": path.stat().st_size,
    }


def extract_notebook(path: Path) -> dict[str, Any]:
    notebook = json.loads(read_text(path))
    cells = []
    for idx, cell in enumerate(notebook.get("cells", [])):
        source = "".join(cell.get("source", [])).strip()
        if source:
            cells.append({
                "index": idx,
                "type": cell.get("cell_type", ""),
                "source": source,
            })
    return {"cell_count": len(notebook.get("cells", [])), "cells": cells}


def gather_runtime_info() -> dict[str, str]:
    source = read_text(ROOT / "main.py")
    result = {
        "input_video_path": "unknown",
        "resize_width": "unknown",
        "max_frames": "unknown",
        "track_stub": "unknown",
        "camera_stub": "unknown",
        "output_path": "unknown",
    }
    for line in source.splitlines():
        stripped = line.strip()
        if stripped.startswith("input_video_path ="):
            result["input_video_path"] = stripped.split("=", 1)[1].strip().strip("'\"")
        elif "read_video(" in stripped:
            if "resize_width=" in stripped:
                result["resize_width"] = stripped.split("resize_width=", 1)[1].split(",", 1)[0].rstrip(")")
            if "max_frames=" in stripped:
                result["max_frames"] = stripped.split("max_frames=", 1)[1].split(")", 1)[0]
        elif "track_stubs" in stripped or "_track_stubs" in stripped:
            if "stub_path=" in stripped:
                result["track_stub"] = stripped.split("stub_path=", 1)[1].rstrip(")")
        elif "camera_movement_stub" in stripped and "stub_path=" in stripped:
            result["camera_stub"] = stripped.split("stub_path=", 1)[1].rstrip(")")
        elif stripped.startswith("save_video("):
            result["output_path"] = stripped.split(",", 1)[1].strip().rstrip(")")
    return result


def dataset_stats() -> dict[str, Any]:
    data_yaml = yaml.safe_load(read_text(ROOT / "training" / "football-players-detection-1" / "data.yaml"))
    dataset_root = ROOT / "training" / "football-players-detection-1" / "football-players-detection-1"
    split_counts = {}
    for split in ("train", "valid", "test"):
        split_counts[split] = {
            "images": len(list((dataset_root / split / "images").glob("*"))),
            "labels": len(list((dataset_root / split / "labels").glob("*"))),
        }
    return {
        "data_yaml": data_yaml,
        "split_counts": split_counts,
        "dataset_readme": read_text(ROOT / "training" / "football-players-detection-1" / "README.dataset.txt").strip(),
        "dataset_roboflow": read_text(ROOT / "training" / "football-players-detection-1" / "README.roboflow.txt").strip(),
    }


def build_context() -> dict[str, Any]:
    py_info = {}
    symbols: list[SymbolInfo] = []
    import_counter = Counter()
    for rel in IMPORTANT_FILES:
        path = ROOT / rel
        if path.suffix == ".py":
            info = extract_py_info(path, rel)
            py_info[rel] = info
            symbols.extend(info["symbols"])
            import_counter.update(info["imports"])

    name_to_symbols = defaultdict(list)
    for symbol in symbols:
        name_to_symbols[symbol.name].append(symbol.qualified_name)
        name_to_symbols[symbol.qualified_name].append(symbol.qualified_name)

    called_by = defaultdict(set)
    for symbol in symbols:
        for call in symbol.calls:
            if call in name_to_symbols:
                for target in name_to_symbols[call]:
                    called_by[target].add(symbol.qualified_name)

    notebooks = {
        "training": extract_notebook(ROOT / "training" / "football_training_yolo_v5.ipynb"),
        "color_assignment": extract_notebook(ROOT / "development_and_analysis" / "color_assignment.ipynb"),
    }

    return {
        "py_info": py_info,
        "symbols": symbols,
        "called_by": {k: sorted(v) for k, v in called_by.items()},
        "runtime": gather_runtime_info(),
        "dataset": dataset_stats(),
        "notebooks": notebooks,
        "requirements": [line.strip() for line in read_text(ROOT / "requirements.txt").splitlines() if line.strip()],
        "tree": list_project_tree(ROOT, depth=3),
        "import_counter": import_counter,
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)

    style_updates = {
        "Title": (24, RGBColor(24, 54, 92)),
        "Heading 1": (16, RGBColor(24, 54, 92)),
        "Heading 2": (12.5, RGBColor(52, 90, 140)),
        "Heading 3": (11, RGBColor(70, 70, 70)),
    }
    for name, (size, color) in style_updates.items():
        style = doc.styles[name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = color


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        set_cell_shading(table.rows[0].cells[i], "DCE6F1")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph("")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def get_font(size: int = 20):
    try:
        return ImageFont.truetype("arial.ttf", size)
    except OSError:
        return ImageFont.load_default()


def draw_box(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], text: str, fill: str, outline: str = "#1A365D") -> None:
    draw.rounded_rectangle(xy, radius=16, fill=fill, outline=outline, width=3)
    font = get_font(20)
    x1, y1, x2, y2 = xy
    wrapped = textwrap.fill(text, width=18)
    bbox = draw.multiline_textbbox((0, 0), wrapped, font=font, spacing=4)
    text_w = bbox[2] - bbox[0]
    text_h = bbox[3] - bbox[1]
    draw.multiline_text(
        (x1 + (x2 - x1 - text_w) / 2, y1 + (y2 - y1 - text_h) / 2),
        wrapped,
        font=font,
        fill="#10233D",
        align="center",
        spacing=4,
    )


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], fill: str = "#355C8C") -> None:
    draw.line([start, end], fill=fill, width=4)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    arrow_size = 10
    p1 = (
        end[0] - arrow_size * math.cos(angle - math.pi / 6),
        end[1] - arrow_size * math.sin(angle - math.pi / 6),
    )
    p2 = (
        end[0] - arrow_size * math.cos(angle + math.pi / 6),
        end[1] - arrow_size * math.sin(angle + math.pi / 6),
    )
    draw.polygon([end, p1, p2], fill=fill)


def create_diagram(path: Path, title: str, boxes: list[tuple[tuple[int, int, int, int], str, str]], arrows: list[tuple[tuple[int, int], tuple[int, int]]]) -> None:
    image = PILImage.new("RGB", (1600, 900), "#F7FAFC")
    draw = ImageDraw.Draw(image)
    title_font = get_font(30)
    draw.text((50, 30), title, fill="#1A365D", font=title_font)
    for xy, text, fill in boxes:
        draw_box(draw, xy, text, fill)
    for start, end in arrows:
        arrow(draw, start, end)
    image.save(path)


def generate_diagrams() -> dict[str, Path]:
    DIAGRAM_DIR.mkdir(exist_ok=True)
    diagrams = {}

    diagrams["architecture"] = DIAGRAM_DIR / "architecture.png"
    create_diagram(
        diagrams["architecture"],
        "System Architecture",
        [
            ((80, 180, 320, 320), "Input Videos", "#DCE6F1"),
            ((420, 180, 700, 320), "main.py Orchestrator", "#E8F0FE"),
            ((800, 120, 1130, 250), "Tracker + YOLO + ByteTrack", "#D9F2E6"),
            ((800, 280, 1130, 410), "Team / Ball / Camera / View / Speed Modules", "#FFF4D6"),
            ((1230, 180, 1510, 320), "Rendered Output Videos", "#FDE7E9"),
        ],
        [
            ((320, 250), (420, 250)),
            ((700, 250), (800, 185)),
            ((700, 250), (800, 345)),
            ((1130, 250), (1230, 250)),
            ((1130, 345), (1230, 250)),
        ],
    )

    diagrams["workflow"] = DIAGRAM_DIR / "workflow.png"
    create_diagram(
        diagrams["workflow"],
        "Runtime Workflow",
        [
            ((60, 140, 290, 250), "Read Video", "#DCE6F1"),
            ((340, 140, 570, 250), "Detect + Track", "#E8F0FE"),
            ((620, 140, 850, 250), "Camera Motion", "#D9F2E6"),
            ((900, 140, 1130, 250), "View Transform", "#FFF4D6"),
            ((1180, 140, 1410, 250), "Speed/Distance", "#FDE7E9"),
            ((340, 360, 570, 470), "Team Assignment", "#E8F0FE"),
            ((620, 360, 850, 470), "Ball Assignment", "#D9F2E6"),
            ((900, 360, 1130, 470), "Draw Overlays", "#FFF4D6"),
            ((1180, 360, 1410, 470), "Write AVI", "#FDE7E9"),
        ],
        [
            ((290, 195), (340, 195)),
            ((570, 195), (620, 195)),
            ((850, 195), (900, 195)),
            ((1130, 195), (1180, 195)),
            ((455, 250), (455, 360)),
            ((735, 250), (735, 360)),
            ((850, 415), (900, 415)),
            ((1130, 415), (1180, 415)),
        ],
    )

    diagrams["training"] = DIAGRAM_DIR / "training_pipeline.png"
    create_diagram(
        diagrams["training"],
        "Training Pipeline",
        [
            ((90, 180, 350, 300), "Roboflow Dataset Download", "#DCE6F1"),
            ((420, 180, 680, 300), "Folder Rearrangement", "#E8F0FE"),
            ((750, 180, 1010, 300), "Ultralytics YOLOv5x Training", "#D9F2E6"),
            ((1080, 180, 1340, 300), "best.pt / last.pt Export", "#FFF4D6"),
        ],
        [
            ((350, 240), (420, 240)),
            ((680, 240), (750, 240)),
            ((1010, 240), (1080, 240)),
        ],
    )

    diagrams["folders"] = DIAGRAM_DIR / "folder_flow.png"
    create_diagram(
        diagrams["folders"],
        "Folder Relationship View",
        [
            ((80, 130, 320, 240), "input_videos", "#DCE6F1"),
            ((80, 310, 320, 420), "models", "#DCE6F1"),
            ((420, 220, 720, 330), "main.py", "#E8F0FE"),
            ((820, 100, 1120, 210), "stubs", "#D9F2E6"),
            ((820, 250, 1120, 360), "trackers / analytics modules", "#FFF4D6"),
            ((820, 400, 1120, 510), "output_videos", "#FDE7E9"),
        ],
        [
            ((320, 185), (420, 250)),
            ((320, 365), (420, 300)),
            ((720, 260), (820, 155)),
            ((720, 275), (820, 305)),
            ((720, 290), (820, 455)),
        ],
    )
    return diagrams


def codeblock(doc: Document, lines: list[str]) -> None:
    para = doc.add_paragraph()
    run = para.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(8.5)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Football Analysis")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(24, 54, 92)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Complete Technical Documentation Report").italic = True

    meta = [
        "Project Type: Computer Vision / Sports Video Analytics / Object Detection and Tracking",
        "Author Name: Not specified in repository",
        "Technologies: Python, Ultralytics YOLO, OpenCV, supervision ByteTrack, NumPy, pandas, scikit-learn",
        "Date: Generated from local repository state",
        "Institution/Organization: Not specified in repository",
    ]
    for line in meta:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)
    doc.add_page_break()


def build_docx(context: dict[str, Any], diagrams: dict[str, Path]) -> None:
    doc = Document()
    style_doc(doc)
    add_title_page(doc)

    doc.add_heading("1. Abstract", level=1)
    doc.add_paragraph(
        "This project builds a football video analysis system that detects the ball, goalkeepers, players, and referees "
        "from video frames, tracks them across time, estimates team identity from jersey color, compensates for camera "
        "movement, projects motion to field coordinates, and renders tactical overlays such as ball possession, player "
        "speed, and cumulative distance. The main objective is to convert raw football footage into structured movement "
        "and possession information that can support analysis, demonstration, or academic project evaluation."
    )

    doc.add_heading("2. Project Overview", level=1)
    add_bullets(doc, [
        "Problem statement: raw sports video is difficult to analyze frame by frame without automated detection, tracking, and summarization.",
        "Primary use case: annotate football clips with player tracking, ball control, and motion-based metrics.",
        "End users: students, evaluators, computer vision practitioners, and developers extending sports analytics projects.",
        "Core features: object detection, object tracking, team assignment, camera motion estimation, perspective mapping, speed estimation, possession display, annotated video export.",
        "Main modules: trackers, team_assigner, player_ball_assigner, camera_movement_estimator, view_transformer, speed_and_distance_estimator, and shared utils.",
    ])
    doc.add_picture(str(diagrams["architecture"]), width=Inches(7.1))
    doc.add_paragraph("")

    doc.add_heading("3. Folder Structure Analysis", level=1)
    doc.add_paragraph("Recursive project tree (limited to major project depth for readability):")
    codeblock(doc, context["tree"])
    doc.add_paragraph("Folder purpose summary:")
    folder_rows = []
    for folder, purpose in sorted(FOLDER_PURPOSES.items()):
        folder_rows.append([folder, purpose])
    add_table(doc, ["Folder", "Purpose"], folder_rows, widths=[2.0, 5.2])
    doc.add_picture(str(diagrams["folders"]), width=Inches(7.1))

    doc.add_heading("4. File-by-File Analysis", level=1)
    for rel in IMPORTANT_FILES:
        path = ROOT / rel
        doc.add_heading(rel, level=2)
        doc.add_paragraph(FILE_PURPOSES.get(rel, "Project file included in the generated report."))
        details = [
            f"File path: {path}",
            f"Exists: {'Yes' if path.exists() else 'No'}",
            f"Size: {path.stat().st_size if path.exists() else 0} bytes",
        ]
        if rel in context["py_info"]:
            info = context["py_info"][rel]
            details.append(f"Approximate line count: {info['line_count']}")
            if info["imports"]:
                details.append("Direct imports: " + "; ".join(info["imports"]))
        add_bullets(doc, details)

        if rel == "main.py":
            add_bullets(doc, [
                f"Current configured input path: {context['runtime']['input_video_path']}",
                f"Current resize_width argument: {context['runtime']['resize_width']}",
                f"Current max_frames argument: {context['runtime']['max_frames']}",
                f"Current track stub pattern: {context['runtime']['track_stub']}",
                f"Current camera stub pattern: {context['runtime']['camera_stub']}",
                f"Current output save path expression: {context['runtime']['output_path']}",
            ])
        elif rel.endswith(".ipynb"):
            nb_key = "training" if "training" in rel else "color_assignment"
            nb = context["notebooks"][nb_key]
            add_bullets(doc, [
                f"Notebook cell count: {nb['cell_count']}",
                "Notebook is used for experimentation and does not participate directly in the main runtime import chain.",
            ])
            preview_rows = []
            for cell in nb["cells"][:10]:
                preview_rows.append([str(cell["index"]), cell["type"], textwrap.shorten(cell["source"].replace("\n", " "), width=90)])
            add_table(doc, ["Cell", "Type", "Role preview"], preview_rows, widths=[0.7, 1.0, 5.5])
        elif rel.endswith("data.yaml"):
            data_yaml = context["dataset"]["data_yaml"]
            add_table(doc, ["Key", "Value"], [
                ["Classes", ", ".join(data_yaml["names"])],
                ["Number of classes", str(data_yaml["nc"])],
                ["Train path", str(data_yaml["train"])],
                ["Validation path", str(data_yaml["val"])],
                ["Test path", str(data_yaml["test"])],
                ["Roboflow URL", str(data_yaml["roboflow"]["url"])],
            ], widths=[2.2, 5.0])

    doc.add_heading("5. Function-by-Function Documentation", level=1)
    symbol_rows = []
    for symbol in context["symbols"]:
        if symbol.kind == "class":
            continue
        called_by = ", ".join(context["called_by"].get(symbol.qualified_name, [])) or "No direct local reference found"
        symbol_rows.append([
            symbol.qualified_name,
            symbol.file,
            ", ".join(symbol.parameters) if symbol.parameters else "-",
            "Not explicitly annotated",
            symbol.kind,
            called_by,
        ])
    add_table(doc, ["Function / Method", "File", "Parameters", "Return Type", "Kind", "Called By"], symbol_rows, widths=[2.0, 1.8, 1.7, 1.1, 0.8, 1.8])

    for rel, info in context["py_info"].items():
        file_symbols = [s for s in context["symbols"] if s.file == rel and s.kind in {"function", "method"}]
        if not file_symbols:
            continue
        doc.add_heading(rel, level=2)
        for symbol in file_symbols:
            doc.add_heading(symbol.qualified_name, level=3)
            doc.add_paragraph(
                f"Parameters: {', '.join(symbol.parameters) if symbol.parameters else 'None declared explicitly beyond defaults.'}"
            )
            add_bullets(doc, [
                f"Symbol kind: {symbol.kind}",
                f"Declared at line: {symbol.line}",
                f"Called by: {', '.join(context['called_by'].get(symbol.qualified_name, [])) or 'No direct local call match found'}",
                f"Internal call names referenced: {', '.join(symbol.calls) if symbol.calls else 'No nested calls detected by AST scan'}",
            ])

    doc.add_heading("6. Notebook Analysis", level=1)
    training_nb = context["notebooks"]["training"]
    doc.add_heading("training/football_training_yolo_v5.ipynb", level=2)
    add_bullets(doc, [
        "Purpose: download the dataset, train the YOLO detector, and export resulting weights.",
        "Cell 0 installs ultralytics and roboflow.",
        "Cell 1 authenticates Roboflow, selects the project/version, and downloads the dataset in YOLOv5 format.",
        "Cell 3 manually moves split folders into the structure expected by the project.",
        "Cell 5 launches training with yolo task=detect mode=train model=yolov5x.pt data={dataset.location}/data.yaml epochs=10 imgsz=640.",
        "Cells 7-10 mount Google Drive and copy best.pt and last.pt into a persistent location.",
    ])
    color_nb = context["notebooks"]["color_assignment"]
    doc.add_heading("development_and_analysis/color_assignment.ipynb", level=2)
    add_bullets(doc, [
        "Purpose: prototype and validate the jersey-color clustering logic later implemented in team_assigner.py.",
        "Loads a cropped player image, converts it to RGB, isolates the top half, reshapes pixels into a 2D feature array, and runs KMeans with 2 clusters.",
        "Uses corner-cluster voting to infer which cluster is background and which one likely belongs to the player's jersey.",
    ])

    doc.add_heading("7. Dataset Analysis", level=1)
    dataset = context["dataset"]
    add_table(doc, ["Property", "Value"], [
        ["Dataset source", dataset["data_yaml"]["roboflow"]["url"]],
        ["License", dataset["data_yaml"]["roboflow"]["license"]],
        ["Format", "YOLOv5 PyTorch detection format"],
        ["Classes", ", ".join(dataset["data_yaml"]["names"])],
        ["Total local image count", str(sum(v["images"] for v in dataset["split_counts"].values()))],
        ["Train images", str(dataset["split_counts"]["train"]["images"])],
        ["Validation images", str(dataset["split_counts"]["valid"]["images"])],
        ["Test images", str(dataset["split_counts"]["test"]["images"])],
        ["Augmentation", "50% horizontal flip and random brightness adjustment between -20% and +20%"],
    ], widths=[2.2, 5.0])
    add_bullets(doc, [
        "Dataset type: image detection dataset, not a tabular CSV dataset.",
        "Label encoding: YOLO bounding-box text files paired with images.",
        "Sample rows / correlation analysis: not applicable in the traditional tabular sense because the project uses image data and YOLO label files.",
        "Image dimensions vary by source clip frame extraction; the training config resizes at training time via imgsz=640.",
    ])

    doc.add_heading("8. Model Architecture Analysis", level=1)
    add_bullets(doc, [
        "Inference model used by the project runtime: models/best.pt loaded through ultralytics.YOLO.",
        "Training base model indicated by notebook: yolov5x.pt.",
        "Model task: object detection with four classes: ball, goalkeeper, player, referee.",
        "Training command in notebook specifies imgsz=640 and epochs=10.",
        "Loss, optimizer, and layer internals are managed inside Ultralytics and are not redefined in project source files.",
        "The repository stores final trained weights rather than a hand-written neural-network class definition.",
    ])
    doc.add_picture(str(diagrams["training"]), width=Inches(7.1))

    doc.add_heading("9. Training Pipeline Explanation", level=1)
    add_bullets(doc, [
        "1. Data loading: notebook downloads the Roboflow dataset with the roboflow SDK.",
        "2. Data preprocessing: folder restructuring and Roboflow-provided augmentation configuration.",
        "3. Feature extraction: handled internally by the YOLO architecture during training.",
        "4. Dataset split usage: train/valid/test paths are defined in data.yaml.",
        "5. Model initialization: yolo detect train starts from yolov5x.pt pretrained weights.",
        "6. Training loop: managed inside the Ultralytics trainer.",
        "7. Loss calculation: managed inside Ultralytics for objectness/classification/box regression.",
        "8. Optimization and backpropagation: handled by Ultralytics runtime.",
        "9. Validation: validation data path is provided in data.yaml and evaluated during training.",
        "10. Model saving: best.pt and last.pt written under runs/detect/train/weights.",
        "11. Export / persistence: notebook copies these weights to Google Drive.",
    ])

    doc.add_heading("10. Input and Output Flow", level=1)
    doc.add_picture(str(diagrams["workflow"]), width=Inches(7.1))
    add_bullets(doc, [
        "User input: a chosen video path in main.py plus model weights in models/best.pt.",
        "Intermediate transformations: frame extraction, detection, tracking, motion adjustment, perspective mapping, team assignment, possession assignment, overlay rendering.",
        "Final output: an AVI video saved to output_videos with per-frame visual analytics overlays.",
        "Cached auxiliary outputs: pickle files in stubs storing tracking dictionaries and camera-motion vectors.",
    ])

    doc.add_heading("11. Library and Dependency Analysis", level=1)
    library_rows = [
        ["ultralytics", "Loads YOLO models, performs inference, and is used by the training notebook for model training."],
        ["supervision", "Provides ByteTrack and detection wrappers for converting Ultralytics results into trackable objects."],
        ["cv2 / OpenCV", "Video I/O, image resizing, drawing, color conversion, optical flow, perspective transform, and text rendering."],
        ["numpy", "Array operations, polygon / point shaping, masking, and trajectory storage."],
        ["pandas", "Interpolates missing ball detections."],
        ["scikit-learn", "Runs KMeans for team color clustering."],
        ["pickle", "Serializes and loads cached stubs."],
        ["yaml", "Loads training metadata from data.yaml."],
        ["roboflow", "Dataset download in the notebook."],
    ]
    add_table(doc, ["Library", "Project role"], library_rows, widths=[1.8, 5.5])

    doc.add_heading("12. Algorithm Explanation", level=1)
    add_table(doc, ["Algorithm", "Where used", "How it works", "Why selected", "Limits"], [
        ["YOLO object detection", "Tracker.detect_frames()", "One-stage detector predicts boxes and classes directly from images.", "Fast and widely used for real-time detection tasks.", "Dependent on training data quality and can miss small fast objects like a football."],
        ["ByteTrack", "Tracker.get_object_tracks()", "Associates detections across frames to preserve ids over time.", "Simple ready-to-use tracker available through supervision.", "Current API is deprecated in installed supervision version."],
        ["KMeans clustering", "TeamAssigner", "Clusters jersey pixel colors into two groups and then clusters players into two teams.", "No manual team labels are needed for the match clip.", "Sensitive to lighting, crop quality, and background contamination."],
        ["Lucas-Kanade optical flow", "CameraMovementEstimator", "Tracks sparse corner features between consecutive frames to estimate camera motion.", "Efficient sparse motion estimation for frame-to-frame compensation.", "Can fail when insufficient stable features are found."],
        ["Perspective transform", "ViewTransformer", "Maps image points inside a four-point polygon into pitch coordinates.", "Provides approximate metric space for speed and distance estimation.", "Requires manually chosen corner coordinates and assumes a fixed geometry."],
        ["Nearest-neighbor possession assignment", "PlayerBallAssigner", "Selects the nearest player foot endpoint to the ball center within a threshold.", "Simple heuristic with low computation cost.", "May be incorrect during close contests, occlusion, or poor ball detection."],
    ], widths=[1.5, 1.7, 2.3, 1.5, 1.3])

    doc.add_heading("13. Evaluation Metrics", level=1)
    add_bullets(doc, [
        "The runtime repository does not compute accuracy, precision, recall, F1-score, ROC-AUC, or confusion matrices directly in code.",
        "README.md references RMSE and MAPE, but those metrics are not implemented in the inspected Python runtime files and appear to be leftover or unrelated documentation content.",
        "Operational runtime quality is mainly observed visually through the correctness of detections, tracks, overlays, and output videos.",
        "Where mathematically relevant, speed is computed as distance / time and converted from m/s to km/h by multiplying by 3.6.",
    ])

    doc.add_heading("14. Visualization and Output Analysis", level=1)
    add_bullets(doc, [
        "Ellipses indicate tracked players and referees. Player ellipses use the assigned team color; referee ellipses are yellow.",
        "Track ids are shown in filled rectangles near the player ellipse base.",
        "A green triangle marks the ball; a red triangle marks the player who currently has the ball according to the possession heuristic.",
        "Ball control percentages are computed cumulatively over processed frames and drawn in the lower-right overlay panel.",
        "Camera movement text is drawn on a translucent overlay near the top-left of the frame.",
        "Speed and distance are drawn below tracked players when transformed positions are available.",
    ])

    doc.add_heading("15. Execution Workflow", level=1)
    codeblock(doc, [
        "pip install -r requirements.txt",
        "python main.py",
        "",
        "# Optional quick detector-only check",
        "python yolo_inference.py",
    ])
    add_bullets(doc, [
        "Place videos in input_videos/.",
        "Place trained best.pt in models/.",
        "If using long videos, adjust resize_width and max_frames in main.py to stay within memory limits.",
        "Per-video stubs should be cleared or renamed when switching inputs if you want fresh recomputation.",
    ])

    doc.add_heading("16. Error Handling and Edge Cases", level=1)
    add_bullets(doc, [
        "Missing supervision package causes ModuleNotFoundError at tracker import time.",
        "Broken Ultralytics roaming-profile settings can cause settings.json errors; current main.py redirects configuration into a local .ultralytics folder.",
        "Using stubs from one video with a different video can cause frame-count mismatch index errors.",
        "If no player is assigned the ball in the earliest frame, possession logic must handle an empty history; current main.py now falls back to 0.",
        "Large videos can exhaust memory because frames are stored in RAM before processing and again during rendering.",
        "Camera motion estimation can fail when no valid corners exist; current implementation now reacquires features and skips bad frames.",
    ])

    doc.add_heading("17. Performance Optimization", level=1)
    add_bullets(doc, [
        "Current optimization strategy: optional stub caching of tracks and camera motion to skip repeated expensive stages.",
        "Current video loading helper supports resize_width and max_frames to reduce memory load on long videos.",
        "Possible improvements: stream frames instead of loading entire videos; write output incrementally; reuse detection tensors; preserve source FPS; avoid duplicate full-frame copies in multiple rendering stages.",
    ])

    doc.add_heading("18. Future Improvements", level=1)
    add_bullets(doc, [
        "Replace or update deprecated ByteTrack integration to the current supervision API.",
        "Add a streaming pipeline that does not keep all frames in memory.",
        "Use stronger ball-tracking or re-identification for more robust possession estimates.",
        "Add evaluation scripts with measurable detection/tracking metrics.",
        "Expose the pipeline as a web app or API for interactive usage.",
        "Parameterize all hard-coded file paths, polygon points, FPS assumptions, and team overrides.",
    ])

    doc.add_heading("19. Conclusion", level=1)
    doc.add_paragraph(
        "The Football_Analysis project successfully combines detector inference, tracking, team clustering, optical-flow based "
        "camera compensation, perspective projection, and trajectory-based metrics into a single end-to-end football analysis "
        "workflow. Its strengths are modularity, readability, and clear extension points. Its main technical limitations are "
        "memory usage, hard-coded assumptions, and partial dependence on notebook/manual setup steps. Even with those limits, "
        "the repository is a strong academic or prototype-level sports analytics implementation."
    )

    doc.save(DOCX_PATH)


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(
        name="ReportTitle",
        parent=styles["Title"],
        fontName="Helvetica-Bold",
        fontSize=20,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#1A365D"),
        spaceAfter=10,
    ))
    styles.add(ParagraphStyle(
        name="Section",
        parent=styles["Heading1"],
        fontName="Helvetica-Bold",
        fontSize=14,
        textColor=colors.HexColor("#1A365D"),
        spaceBefore=10,
        spaceAfter=6,
    ))
    styles.add(ParagraphStyle(
        name="SubSection",
        parent=styles["Heading2"],
        fontName="Helvetica-Bold",
        fontSize=10.5,
        textColor=colors.HexColor("#355C8C"),
        spaceBefore=6,
        spaceAfter=4,
    ))
    styles["BodyText"].fontName = "Helvetica"
    styles["BodyText"].fontSize = 9
    styles["BodyText"].leading = 11
    return styles


def pdf_table(data: list[list[str]], widths: list[float]) -> Table:
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#DCE6F1")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1A365D")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LEFTPADDING", (0, 0), (-1, -1), 4),
        ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    return table


def build_pdf(context: dict[str, Any], diagrams: dict[str, Path]) -> None:
    styles = pdf_styles()
    story = []
    story.append(Paragraph("Football Analysis Full Technical Report", styles["ReportTitle"]))
    story.append(Paragraph("Detailed project documentation generated from the local repository state.", styles["BodyText"]))
    story.append(Spacer(1, 0.3 * cm))

    sections = [
        ("Abstract", [
            "This project analyzes football videos with a pipeline that detects objects, tracks them across frames, assigns teams, estimates motion, maps coordinates to a pitch, and renders analytics overlays.",
        ]),
        ("Project Overview", [
            "Main goal: transform raw football footage into structured visual analytics outputs.",
            "Core features: detection, tracking, team color clustering, camera motion estimation, perspective transform, ball assignment, speed and distance overlays.",
        ]),
        ("Runtime Configuration", [
            f"Input video path: {context['runtime']['input_video_path']}",
            f"resize_width: {context['runtime']['resize_width']}",
            f"max_frames: {context['runtime']['max_frames']}",
            f"Output save path expression: {context['runtime']['output_path']}",
        ]),
        ("Dataset Summary", [
            f"Dataset URL: {context['dataset']['data_yaml']['roboflow']['url']}",
            f"Classes: {', '.join(context['dataset']['data_yaml']['names'])}",
            f"Split counts: train={context['dataset']['split_counts']['train']['images']}, valid={context['dataset']['split_counts']['valid']['images']}, test={context['dataset']['split_counts']['test']['images']}",
        ]),
    ]
    for title, bullets in sections:
        story.append(Paragraph(title, styles["Section"]))
        for bullet in bullets:
            story.append(Paragraph(f"- {bullet}", styles["BodyText"]))
        story.append(Spacer(1, 0.15 * cm))

    story.append(Image(str(diagrams["architecture"]), width=16 * cm, height=9 * cm))
    story.append(Spacer(1, 0.2 * cm))
    story.append(Image(str(diagrams["workflow"]), width=16 * cm, height=9 * cm))
    story.append(PageBreak())

    story.append(Paragraph("Important File Inventory", styles["Section"]))
    file_rows = [["File", "Purpose"]]
    for rel in IMPORTANT_FILES:
        file_rows.append([rel, FILE_PURPOSES.get(rel, "Project file")])
    story.append(pdf_table(file_rows, [6.0 * cm, 11.2 * cm]))

    story.append(PageBreak())
    story.append(Paragraph("Function and Method Inventory", styles["Section"]))
    symbol_rows = [["Symbol", "File", "Parameters", "Called By"]]
    for symbol in context["symbols"]:
        if symbol.kind == "class":
            continue
        symbol_rows.append([
            symbol.qualified_name,
            symbol.file,
            ", ".join(symbol.parameters) if symbol.parameters else "-",
            ", ".join(context["called_by"].get(symbol.qualified_name, [])) or "No direct local match",
        ])
    story.append(pdf_table(symbol_rows, [4.5 * cm, 4.0 * cm, 4.0 * cm, 5.0 * cm]))

    story.append(PageBreak())
    story.append(Paragraph("Libraries and Dependencies", styles["Section"]))
    dep_rows = [["Dependency", "Where/Why used"]]
    dep_rows.extend([
        ["ultralytics", "Detection inference and notebook-based training"],
        ["supervision", "ByteTrack and detection conversion"],
        ["opencv-python", "Video I/O, drawing, optical flow, transforms"],
        ["numpy", "Arrays and geometry"],
        ["pandas", "Ball interpolation"],
        ["scikit-learn", "KMeans team assignment"],
    ])
    story.append(pdf_table(dep_rows, [4.5 * cm, 12.8 * cm]))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=A4,
        leftMargin=1.2 * cm,
        rightMargin=1.2 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm,
    )
    doc.build(story)


def write_summary_files(context: dict[str, Any]) -> None:
    SUMMARY_PATH.write_text(
        "\n".join([
            "# Football Analysis Presentation Points",
            "",
            "## 1. Problem",
            "- Raw football videos are difficult to analyze manually frame by frame.",
            "",
            "## 2. Solution",
            "- This project automates detection, tracking, team assignment, possession estimation, and speed/distance overlays.",
            "",
            "## 3. Core Pipeline",
            "- Read video",
            "- Detect objects with YOLO",
            "- Track with ByteTrack",
            "- Estimate camera motion",
            "- Transform positions to pitch coordinates",
            "- Compute speed, distance, and ball control",
            "- Render annotated output video",
            "",
            "## 4. Dataset",
            f"- Roboflow dataset URL: {context['dataset']['data_yaml']['roboflow']['url']}",
            f"- Classes: {', '.join(context['dataset']['data_yaml']['names'])}",
            "",
            "## 5. Technical Strengths",
            "- Modular pipeline",
            "- Uses standard CV/ML building blocks",
            "- Supports per-video caching through stub pickle files",
            "",
            "## 6. Current Limitations",
            "- Memory-heavy on long videos",
            "- Some hard-coded assumptions remain",
            "- ByteTrack API warning in current supervision version",
        ]),
        encoding="utf-8",
    )

    HIGHLIGHTS_PATH.write_text(
        "\n".join([
            "# Football Analysis Technical Highlights",
            "",
            "- End-to-end sports analytics video pipeline implemented in Python.",
            "- Custom-trained YOLO detector saved as models/best.pt.",
            "- Team assignment built from KMeans clustering on jersey color crops.",
            "- Camera motion compensation implemented with Lucas-Kanade optical flow.",
            "- Perspective transform maps player movement into approximate field coordinates.",
            "- Speed and cumulative distance are rendered directly onto output frames.",
            "- Per-video stub caching reduces repeated detector and motion computation.",
        ]),
        encoding="utf-8",
    )


def main() -> None:
    REPORT_DIR.mkdir(exist_ok=True)
    context = build_context()
    diagrams = generate_diagrams()
    build_docx(context, diagrams)
    build_pdf(context, diagrams)
    write_summary_files(context)
    print(f"Created {DOCX_PATH}")
    print(f"Created {PDF_PATH}")
    for name, path in diagrams.items():
        print(f"Created diagram {name}: {path}")
    print(f"Created {SUMMARY_PATH}")
    print(f"Created {HIGHLIGHTS_PATH}")


if __name__ == "__main__":
    main()
